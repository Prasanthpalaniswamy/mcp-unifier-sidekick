import json
import mimetypes
import os
import smtplib
import zipfile
from datetime import datetime
from email.message import EmailMessage
from typing import Any


def _normalize_to_records(content: Any) -> list[dict[str, Any]]:
    """
    Convert supported content into a list of flat records for CSV/XLSX export.

    Supported input:
    - list[dict]
    - dict containing a `data` key with list[dict]
    - dict
    - JSON string representing any of the above
    """
    if isinstance(content, str):
        try:
            content = json.loads(content)
        except json.JSONDecodeError as exc:
            raise ValueError("content must be valid JSON when provided as a string") from exc

    if isinstance(content, dict) and isinstance(content.get("data"), list):
        records = content["data"]
    elif isinstance(content, list):
        records = content
    elif isinstance(content, dict):
        records = [content]
    else:
        raise ValueError(
            "Unsupported content format. Provide a dict, list of dicts, or JSON string."
        )

    if not records:
        return []

    normalized = []
    for item in records:
        if isinstance(item, dict):
            normalized.append(item)
        else:
            normalized.append({"value": item})

    return normalized


def export_content_to_file(
    content: Any,
    output_file: str,
    sheet_name: str = "Sheet1",
) -> str:
    """
    Export arbitrary fetched content to .csv or .xlsx.

    `output_file` may be relative or absolute.
    Supported extensions: .csv, .xlsx
    """
    import pandas as pd

    records = _normalize_to_records(content)
    if not records:
        raise ValueError("No records available to export.")

    df = pd.json_normalize(records)

    output_file = os.path.normpath(output_file)
    extension = os.path.splitext(output_file)[1].lower()

    if extension not in {".csv", ".xlsx"}:
        raise ValueError("output_file must end with .csv or .xlsx")

    output_dir = os.path.dirname(output_file)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    try:
        if extension == ".csv":
            df.to_csv(output_file, index=False, encoding="utf-8-sig")
        else:
            df.to_excel(output_file, index=False, sheet_name=sheet_name)
        return f"Exported {len(df)} row(s) to: {output_file}"
    except PermissionError:
        base, ext = os.path.splitext(output_file)
        fallback = f"{base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
        if extension == ".csv":
            df.to_csv(fallback, index=False, encoding="utf-8-sig")
        else:
            df.to_excel(fallback, index=False, sheet_name=sheet_name)
        return f"Target file was locked. Exported {len(df)} row(s) to fallback file: {fallback}"


def _escape_pdf_text(value: Any) -> str:
    return str(value).replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _build_simple_pdf(lines: list[str]) -> bytes:
    content_parts = ["BT", "/F1 10 Tf", "50 780 Td", "14 TL"]
    first_line = True
    for line in lines:
        safe_line = _escape_pdf_text(line)
        if first_line:
            content_parts.append(f"({safe_line}) Tj")
            first_line = False
        else:
            content_parts.append("T*")
            content_parts.append(f"({safe_line}) Tj")
    content_parts.append("ET")
    stream_text = "\n".join(content_parts) + "\n"
    stream_bytes = stream_text.encode("latin-1", errors="replace")

    objects = []
    objects.append(b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n")
    objects.append(b"2 0 obj << /Type /Pages /Count 1 /Kids [3 0 R] >> endobj\n")
    objects.append(
        b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n"
    )
    objects.append(b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n")
    objects.append(
        f"5 0 obj << /Length {len(stream_bytes)} >> stream\n".encode("latin-1")
        + stream_bytes
        + b"endstream\nendobj\n"
    )

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj)

    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(offsets)}\n".encode("latin-1"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))

    pdf.extend(
        (
            f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF"
        ).encode("latin-1")
    )
    return bytes(pdf)


def export_content_to_pdf(content: Any, output_file: str) -> str:
    """
    Export fetched content to a simple PDF table-like text report.

    Supported input is the same as export_content_to_file.
    """
    import pandas as pd

    records = _normalize_to_records(content)
    if not records:
        raise ValueError("No records available to export.")

    df = pd.json_normalize(records).fillna("")
    output_file = os.path.normpath(output_file)

    if os.path.splitext(output_file)[1].lower() != ".pdf":
        raise ValueError("output_file must end with .pdf")

    output_dir = os.path.dirname(output_file)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    lines = [f"Exported Data Report ({len(df)} row(s))", "-" * 80]
    for row_index, row in enumerate(df.to_dict(orient="records"), start=1):
        lines.append(f"Row {row_index}")
        for key, value in row.items():
            text = str(value)
            while len(text) > 90:
                lines.append(f"  {key}: {text[:90]}")
                text = text[90:]
                key = "..."
            lines.append(f"  {key}: {text}")
        lines.append("-" * 80)

    pdf_bytes = _build_simple_pdf(lines)

    try:
        with open(output_file, "wb") as pdf_file:
            pdf_file.write(pdf_bytes)
        return f"Exported {len(df)} row(s) to: {output_file}"
    except PermissionError:
        base, ext = os.path.splitext(output_file)
        fallback = f"{base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
        with open(fallback, "wb") as pdf_file:
            pdf_file.write(pdf_bytes)
        return f"Target file was locked. Exported {len(df)} row(s) to fallback file: {fallback}"


def export_content_to_docx(content: Any, output_file: str) -> str:
    """
    Export fetched content to a .docx file.

    Supported input is the same as export_content_to_file.
    """
    from docx import Document
    import pandas as pd

    records = _normalize_to_records(content)
    if not records:
        raise ValueError("No records available to export.")

    df = pd.json_normalize(records).fillna("")
    output_file = os.path.normpath(output_file)

    if os.path.splitext(output_file)[1].lower() != ".docx":
        raise ValueError("output_file must end with .docx")

    output_dir = os.path.dirname(output_file)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    document = Document()
    document.add_heading("Exported Data Report", level=1)
    document.add_paragraph(f"Rows exported: {len(df)}")

    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, column in enumerate(df.columns):
        header_cells[index].text = str(column)

    for row in df.to_dict(orient="records"):
        row_cells = table.add_row().cells
        for index, column in enumerate(df.columns):
            row_cells[index].text = str(row.get(column, ""))

    try:
        document.save(output_file)
        return f"Exported {len(df)} row(s) to: {output_file}"
    except PermissionError:
        base, ext = os.path.splitext(output_file)
        fallback = f"{base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
        document.save(fallback)
        return f"Target file was locked. Exported {len(df)} row(s) to fallback file: {fallback}"


def convert_docx_to_pdf_content_based(input_docx_file: str, output_pdf_file: str) -> str:
    """
    Convert a DOCX file to a PDF by extracting paragraphs and tables,
    then regenerating a simple PDF representation.

    This is a content-based conversion, not a layout-faithful rendering.
    """
    from docx import Document

    input_docx_file = os.path.normpath(input_docx_file)
    output_pdf_file = os.path.normpath(output_pdf_file)

    if not os.path.exists(input_docx_file):
        raise FileNotFoundError(f"Input DOCX file not found: {input_docx_file}")
    if os.path.splitext(input_docx_file)[1].lower() != ".docx":
        raise ValueError("input_docx_file must end with .docx")
    if os.path.splitext(output_pdf_file)[1].lower() != ".pdf":
        raise ValueError("output_pdf_file must end with .pdf")

    output_dir = os.path.dirname(output_pdf_file)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    document = Document(input_docx_file)
    lines = [f"Converted from DOCX: {os.path.basename(input_docx_file)}", "=" * 80]

    paragraph_count = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
            paragraph_count += 1

    table_count = 0
    for table in document.tables:
        table_count += 1
        lines.append("-" * 80)
        lines.append(f"Table {table_count}")
        for row in table.rows:
            cell_values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append(" | ".join(cell_values))

    if paragraph_count == 0 and table_count == 0:
        lines.append("No readable paragraphs or tables were found in the DOCX file.")

    pdf_bytes = _build_simple_pdf(lines)

    try:
        with open(output_pdf_file, "wb") as pdf_file:
            pdf_file.write(pdf_bytes)
        return f"Converted DOCX content to PDF: {output_pdf_file}"
    except PermissionError:
        base, ext = os.path.splitext(output_pdf_file)
        fallback = f"{base}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}"
        with open(fallback, "wb") as pdf_file:
            pdf_file.write(pdf_bytes)
        return f"Target file was locked. Converted DOCX content to fallback PDF: {fallback}"


def send_email_via_smtp(
    smtp_host: str,
    smtp_port: int,
    username: str,
    password: str,
    sender_email: str,
    to_emails: str,
    subject: str,
    body: str,
    attachment_paths: list[str] | None = None,
    use_starttls: bool = True,
) -> str:
    """
    Send a simple plain-text email via SMTP.

    to_emails: comma-separated recipient list.
    """
    recipients = [email.strip() for email in to_emails.split(",") if email.strip()]
    if not recipients:
        raise ValueError("At least one recipient email is required.")

    message = EmailMessage()
    message["From"] = sender_email
    message["To"] = ", ".join(recipients)
    message["Subject"] = subject
    message.set_content(body)

    for attachment_path in attachment_paths or []:
        normalized_path = os.path.normpath(attachment_path)
        with open(normalized_path, "rb") as attachment_file:
            data = attachment_file.read()

        mime_type, _ = mimetypes.guess_type(normalized_path)
        if mime_type:
            maintype, subtype = mime_type.split("/", 1)
        else:
            maintype, subtype = "application", "octet-stream"

        message.add_attachment(
            data,
            maintype=maintype,
            subtype=subtype,
            filename=os.path.basename(normalized_path),
        )

    with smtplib.SMTP(smtp_host, smtp_port, timeout=60) as server:
        server.ehlo()
        if use_starttls:
            server.starttls()
            server.ehlo()
        server.login(username, password)
        server.send_message(message)

    return f"Email sent successfully to: {', '.join(recipients)}"


def compress_files_to_zip(
    input_paths: list[str],
    output_zip_file: str,
) -> str:
    """
    Compress one or more files into a ZIP archive.
    """
    normalized_inputs = [os.path.normpath(path) for path in input_paths if path.strip()]
    if not normalized_inputs:
        raise ValueError("At least one input file path is required.")

    output_zip_file = os.path.normpath(output_zip_file)
    if os.path.splitext(output_zip_file)[1].lower() != ".zip":
        raise ValueError("output_zip_file must end with .zip")

    output_dir = os.path.dirname(output_zip_file)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    with zipfile.ZipFile(output_zip_file, "w", compression=zipfile.ZIP_DEFLATED) as zip_file:
        for input_path in normalized_inputs:
            if not os.path.exists(input_path):
                raise FileNotFoundError(f"Input file not found: {input_path}")
            zip_file.write(input_path, arcname=os.path.basename(input_path))

    size_bytes = os.path.getsize(output_zip_file)
    return f"Created ZIP archive: {output_zip_file} ({size_bytes} bytes)"


def split_file(
    input_file: str,
    output_dir: str | None = None,
    part_size_mb: int = 10,
) -> dict[str, Any]:
    """
    Split a file into sequential .partNNN chunks.
    """
    input_file = os.path.normpath(input_file)
    if not os.path.exists(input_file):
        raise FileNotFoundError(f"Input file not found: {input_file}")

    if part_size_mb <= 0:
        raise ValueError("part_size_mb must be greater than 0")

    if output_dir is None or not output_dir.strip():
        output_dir = os.path.dirname(input_file) or "."
    output_dir = os.path.normpath(output_dir)
    os.makedirs(output_dir, exist_ok=True)

    part_size_bytes = part_size_mb * 1024 * 1024
    base_name = os.path.basename(input_file)
    part_files = []

    with open(input_file, "rb") as source:
        part_number = 1
        while True:
            chunk = source.read(part_size_bytes)
            if not chunk:
                break

            part_name = f"{base_name}.part{part_number:03d}"
            part_path = os.path.join(output_dir, part_name)
            with open(part_path, "wb") as part_file:
                part_file.write(chunk)
            part_files.append(part_path)
            part_number += 1

    return {
        "input_file": input_file,
        "output_dir": output_dir,
        "part_size_mb": part_size_mb,
        "parts": part_files,
        "part_count": len(part_files),
    }


def create_reassembly_instructions(
    original_file_name: str,
    part_files: list[str],
) -> str:
    """
    Build human-readable Windows reassembly instructions for split files.
    """
    file_names = [os.path.basename(path) for path in part_files]
    joined_parts = "+".join(file_names)
    return (
        f"To recreate {original_file_name} on Windows, save all parts in one folder and run:\n"
        f"copy /b {joined_parts} {original_file_name}\n"
        f"After the file is rebuilt, open the ZIP normally."
    )